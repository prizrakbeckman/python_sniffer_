import docx,os

def useDocFile(file):
    
    doc = docx.Document(file)
    i = len(doc.paragraphs)
    print(i)
    print(doc.paragraphs[0].text)
    #Showing content of a doc file
    for par in doc.paragraphs:
       print(par.text)
    print('Enter an integer')
    h = input()
    if h.isdigit():
        print('Showing content by smaller parts')
        for par in doc.paragraphs:
            for run in par.runs:
                print('************************')
                print(run.text)
    else:
        print('Showing shit')

def getFullText(filename):
    
    doc = docx.Document(filename)
    fullText = []
    print('************************')
    print('Showing whole text')
    for par in doc.paragraphs:
        fullText.append(par.text)
    return '\n'.join(fullText)


def writeNewWordDocumentFromSomething(filename,pics):
    
    doc = docx.Document(filename)
    doc1 = docx.Document()
    doc1.add_heading('Header 0', 0)
    doc1.add_paragraph('First line of Zbi')
    doc1.add_paragraph('************************************')
    doc1.add_paragraph('\n')
    for par in doc.paragraphs:
        doc1.add_paragraph(par.text)
    print('************************************')
    doc1.add_paragraph('************************************')
    doc1.add_paragraph('End of line of Zbi')
    for pic in pics:
        print(pic)
        par = doc1.add_paragraph()
        run = par.add_run()
        #doc.add_picture(os.path.realpath(pic), width=docx.shared.Inches(1),height=docx.shared.Cm(4))
        run.add_picture(pic)
        run.add_text('\nRoger and Rayleigh')
        print('pics being added')
    doc1.save('helloWorldOfZbi8.docx')
    
os.chdir('C:\\Users\\HP\\Desktop\\titiza')
#useDocFile('VPNApplication.docx')
#print(getFullText('VPNApplication.docx'))
#print('All text of '+os.path.realpath('VPNApplication.docx')+' shown')
pics = []
pics.append('rogerRayleigh.jpg')
writeNewWordDocumentFromSomething('VPNApplication.docx',pics)

